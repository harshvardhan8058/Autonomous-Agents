"""DOCX Generator — renders the executed plan into a professional document."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.schemas.models import Plan, PlannedTask, TaskExecution
from app.tools.dates import timestamp_display
from app.tools.timeline import build_timeline
from app.tools.titles import document_title

_ACCENT = RGBColor(0x1F, 0x3A, 0x5F)
_MUTED = RGBColor(0x6B, 0x72, 0x80)
_HEADER_BG = "1F3A5F"  # hex string for XML shading (no #)


# ---------------------------------------------------------------------------
# Public entry-point
# ---------------------------------------------------------------------------

def generate_docx(
    plan: Plan, executions: list[TaskExecution], output_path: Path
) -> None:
    doc = Document()
    _style_base(doc)

    _title_page(doc, plan)
    _table_of_contents(doc, plan, executions)
    _executive_summary(doc, plan)
    _timeline_table(doc, plan)
    _sections(doc, executions)
    _footer(doc)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


# ---------------------------------------------------------------------------
# Styles
# ---------------------------------------------------------------------------

def _style_base(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    for level, size in (("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)):
        style = doc.styles[level]
        style.font.color.rgb = _ACCENT
        style.font.size = Pt(size)
        style.font.bold = True


# ---------------------------------------------------------------------------
# Title page
# ---------------------------------------------------------------------------

def _title_page(doc: Document, plan: Plan) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(document_title(plan.goal))
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = _ACCENT

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(f"Generated {timestamp_display()}")
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = _MUTED

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Table of Contents
# ---------------------------------------------------------------------------

def _table_of_contents(
    doc: Document, plan: Plan, executions: list[TaskExecution]
) -> None:
    """Build a TOC page that is immediately readable AND auto-updates in Word.

    Strategy:
    - Emit a standard Word complex-field TOC (fldChar begin/separate/end) so
      Word/LibreOffice auto-refreshes it on open.
    - Between 'separate' and 'end' inject one styled paragraph per known
      heading so users see real content without needing to click Update Field.
    - Also set updateFields=true in document settings as a belt-and-braces.
    """
    doc.add_heading("Table of Contents", level=1)

    # ---- collect the headings we are about to write ----
    toc_entries: list[str] = [
        "Executive Summary",
        "Delivery Timeline",
    ]
    for ex in executions:
        if ex.section is not None:
            toc_entries.append(ex.section.heading)

    # ---- helper XML builders ----
    def _fld_char(fld_type: str) -> OxmlElement:
        fc = OxmlElement("w:fldChar")
        fc.set(qn("w:fldCharType"), fld_type)
        return fc

    def _instr(text: str) -> OxmlElement:
        el = OxmlElement("w:instrText")
        el.set(qn("xml:space"), "preserve")
        el.text = text
        return el

    # ---- paragraph 1: fldChar begin + instrText ----
    p_begin = doc.add_paragraph()
    r_begin = p_begin.add_run()
    r_begin._r.append(_fld_char("begin"))
    r_begin._r.append(_instr(' TOC \\o "1-3" \\h \\z \\u '))

    # ---- paragraph 2: fldChar separate (starts the display area) ----
    p_sep = doc.add_paragraph()
    r_sep = p_sep.add_run()
    r_sep._r.append(_fld_char("separate"))

    # ---- real visible TOC entries between separate and end ----
    for i, heading in enumerate(toc_entries):
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Inches(0.2 if i >= 2 else 0)
        # heading text
        r_text = p.add_run(heading)
        r_text.font.name = "Calibri"
        r_text.font.size = Pt(11)
        r_text.font.color.rgb = _ACCENT
        # tab leader dot ....... then page number placeholder
        p.add_run("\t")
        r_pg = p.add_run("—")
        r_pg.font.color.rgb = _MUTED
        r_pg.font.size = Pt(10)
        # set tab stop at right margin for page numbers
        pPr = p._p.get_or_add_pPr()
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:leader"), "dot")
        tab.set(qn("w:pos"), "9000")
        tabs.append(tab)
        pPr.append(tabs)

    # ---- paragraph 3: fldChar end ----
    p_end = doc.add_paragraph()
    r_end = p_end.add_run()
    r_end._r.append(_fld_char("end"))

    # ---- enable auto-update on open ----
    settings = doc.settings.element
    # avoid duplicate updateFields elements
    for existing in settings.findall(qn("w:updateFields")):
        settings.remove(existing)
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Executive summary
# ---------------------------------------------------------------------------

def _executive_summary(doc: Document, plan: Plan) -> None:
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(plan.goal)
    if plan.assumptions:
        h = doc.add_paragraph("Key assumptions:")
        h.runs[0].bold = True
        for assumption in plan.assumptions:
            doc.add_paragraph(assumption, style="List Bullet")


# ---------------------------------------------------------------------------
# Delivery timeline table
# ---------------------------------------------------------------------------

def _shade_cell(cell, hex_color: str) -> None:
    """Apply solid background shading to a table cell via raw XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_col_width(table, col_idx: int, width_inches: float) -> None:
    for row in table.rows:
        row.cells[col_idx].width = Inches(width_inches)


_PRIORITY_COLOR = {
    "High": RGBColor(0xC0, 0x39, 0x2B),
    "Medium": RGBColor(0xE6, 0x7E, 0x22),
    "Low": RGBColor(0x27, 0xAE, 0x60),
}


def _timeline_table(doc: Document, plan: Plan) -> None:
    doc.add_heading("Delivery Timeline", level=1)
    rows = build_timeline(plan.tasks)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"  # always present in python-docx

    # Header row
    hdr = table.rows[0]
    for i, label in enumerate(("Phase", "Priority", "Target")):
        cell = hdr.cells[i]
        _shade_cell(cell, _HEADER_BG)
        cell.text = label
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(11)

    # Data rows
    for row in rows:
        cells = table.add_row().cells
        cells[0].text = row["phase"]

        priority_text = row["priority"]
        p_run = cells[1].paragraphs[0].add_run(priority_text)
        p_run.font.bold = True
        p_run.font.color.rgb = _PRIORITY_COLOR.get(priority_text, _ACCENT)
        cells[1].paragraphs[0].runs[0].text = ""  # clear default text
        cells[1].paragraphs[0].runs[0].font.bold = True
        cells[1].paragraphs[0].runs[0].font.color.rgb = _PRIORITY_COLOR.get(
            priority_text, _ACCENT
        )
        cells[1].text = ""
        r = cells[1].paragraphs[0].add_run(priority_text)
        r.font.bold = True
        r.font.color.rgb = _PRIORITY_COLOR.get(priority_text, _ACCENT)

        cells[2].text = row["target"]
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Column widths: Phase wide, Priority narrow, Target medium
    _set_col_width(table, 0, 3.2)
    _set_col_width(table, 1, 1.2)
    _set_col_width(table, 2, 1.8)


# ---------------------------------------------------------------------------
# Content sections
# ---------------------------------------------------------------------------

def _sections(doc: Document, executions: list[TaskExecution]) -> None:
    for execution in executions:
        if execution.section is None:
            continue
        doc.add_heading(execution.section.heading, level=1)
        for paragraph in execution.section.paragraphs:
            doc.add_paragraph(paragraph)
        for bullet in execution.section.bullets:
            doc.add_paragraph(bullet, style="List Bullet")


# ---------------------------------------------------------------------------
# Footer
# ---------------------------------------------------------------------------

def _footer(doc: Document) -> None:
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(
        f"Autonomous AI Agent  ·  Generated {timestamp_display()}  ·  Confidential"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = _MUTED
